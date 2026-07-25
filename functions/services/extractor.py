"""
Climactix AI — File Extraction Service
Supports: PDF (PyMuPDF), DOCX (python-docx), XLSX (openpyxl)
"""

import io
import re


def extract_text(content: bytes, filename: str, content_type: str = "") -> str:
    """Route file to the correct extractor based on extension / MIME type."""
    fn = filename.lower()

    if fn.endswith(".pdf") or "pdf" in content_type:
        return _extract_pdf(content)
    elif fn.endswith(".docx") or "word" in content_type or "openxmlformats" in content_type:
        return _extract_docx(content)
    elif fn.endswith(".xlsx") or "spreadsheet" in content_type or "excel" in content_type:
        return _extract_xlsx(content)
    elif fn.endswith(".csv") or "text/csv" in content_type or "csv" in content_type:
        return _extract_csv(content)
    elif fn.endswith(".txt") or "text/plain" in content_type:
        return content.decode("utf-8", errors="replace")
    else:
        # Best-effort: try UTF-8 text decode
        try:
            return content.decode("utf-8", errors="replace")
        except Exception:
            raise ValueError(f"Unsupported file type: {filename}")


def _extract_pdf(content: bytes) -> str:
    """Extract text from PDF using PyMuPDF."""
    try:
        import fitz  # PyMuPDF

        parts = []
        with fitz.open(stream=content, filetype="pdf") as doc:
            for i, page in enumerate(doc):
                text = page.get_text("text")
                if text.strip():
                    parts.append(f"[Page {i + 1}]\n{text.strip()}")

        full_text = "\n\n".join(parts)
        return _clean_text(full_text)
    except ImportError:
        raise RuntimeError("PyMuPDF not installed. Run: pip install pymupdf")


def _extract_docx(content: bytes) -> str:
    """Extract text from DOCX using python-docx.

    Each paragraph/table-row is prefixed with a locator (¶N / TableT.RowR)
    so an AI review's citation can point at a real position in the source
    document instead of just "somewhere in this file" — needed for the
    Question Intelligence drawer's "View Source" feature.
    """
    try:
        from docx import Document

        doc = Document(io.BytesIO(content))
        lines = []

        para_idx = 0
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                para_idx += 1
                lines.append(f"[¶{para_idx}] {text}")

        for table_idx, table in enumerate(doc.tables, start=1):
            for row_idx, row in enumerate(table.rows, start=1):
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    lines.append(f"[Table {table_idx}.Row {row_idx}] " + " | ".join(cells))

        return _clean_text("\n\n".join(lines))
    except ImportError:
        raise RuntimeError("python-docx not installed. Run: pip install python-docx")


def _extract_xlsx(content: bytes) -> str:
    """Extract text from XLSX using openpyxl.

    Each row keeps its real 1-based spreadsheet row number as a locator
    (e.g. "[Sheet: Operating Assets, Row 12]") so citations can point at
    an actual row instead of just the filename.
    """
    try:
        import openpyxl

        wb = openpyxl.load_workbook(io.BytesIO(content), data_only=True)
        parts = []

        for sheet in wb.worksheets:
            parts.append(f"=== Sheet: {sheet.title} ===")
            for row_idx, row in enumerate(sheet.iter_rows(values_only=True), start=1):
                cells = [str(c) for c in row if c is not None and str(c).strip() not in ("", "None")]
                if cells:
                    parts.append(f"[Sheet: {sheet.title}, Row {row_idx}] " + "  |  ".join(cells))

        return _clean_text("\n".join(parts))
    except ImportError:
        raise RuntimeError("openpyxl not installed. Run: pip install openpyxl")


def _extract_csv(content: bytes) -> str:
    """Extract text from CSV — renders as labelled rows for the AI.

    Rows keep their real 1-based CSV row number (header = row 1) as a
    locator, matching how a human would refer to "row 12" when opening
    the same file in a spreadsheet app.
    """
    import csv

    try:
        text = content.decode("utf-8", errors="replace")
        reader = csv.reader(text.splitlines())
        rows   = list(reader)
        if not rows:
            return ""

        headers = rows[0] if rows else []
        parts   = []

        for row_idx, row in enumerate(rows[1:], start=2):
            if not any(cell.strip() for cell in row):
                continue
            # Pair each cell with its header for readable context
            pairs = []
            for i, cell in enumerate(row):
                if cell.strip():
                    label = headers[i].strip() if i < len(headers) else f"Column {i+1}"
                    pairs.append(f"{label}: {cell.strip()}")
            if pairs:
                parts.append(f"[Row {row_idx}] " + " | ".join(pairs))

        return _clean_text("\n".join(parts))
    except Exception as e:
        raise RuntimeError(f"CSV extraction failed: {e}")


def _clean_text(text: str) -> str:
    """Remove excessive whitespace and normalize text."""
    # Collapse 3+ blank lines to 2
    text = re.sub(r"\n{3,}", "\n\n", text)
    # Remove non-printable characters except newlines/tabs
    text = re.sub(r"[^\x09\x0A\x0D\x20-\x7E\u00A0-\uFFFF]", " ", text)
    return text.strip()
