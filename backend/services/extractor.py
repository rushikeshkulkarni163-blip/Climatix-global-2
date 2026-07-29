"""
Climactix AI — File Extraction Service
Supports: PDF (PyMuPDF, with OCR fallback for scanned pages), DOCX (python-docx),
XLSX (openpyxl), CSV, images/scanned documents (Pillow + pytesseract), and web URLs
(httpx + BeautifulSoup).

Every extractor tags its output with a real locator (page/paragraph/row/section)
so a downstream AI review's citation can point at an actual position in the
source rather than "somewhere in this file" (Question Intelligence drawer /
Evidence Intelligence Agent "View Source" feature — CLAUDE.md traceability rule).
OCR output is tagged distinctly ("(OCR)") so a lower-confidence machine-read
page is never presented as indistinguishable from a native text layer.
"""

import io
import re

# Pages with fewer non-whitespace characters than this are treated as
# "no real text layer" and sent through OCR instead — a native text layer
# almost never renders this sparse, but a scanned page with a stray header
# might, so this is a floor, not a hard proof of scan-vs-native.
_OCR_FALLBACK_CHAR_THRESHOLD = 20


def extract_text(content: bytes, filename: str, content_type: str = "") -> str:
    """Route file to the correct extractor based on extension / MIME type."""
    fn = filename.lower()

    if fn.endswith(".pdf") or "pdf" in content_type:
        return _extract_pdf(content)
    elif fn.endswith(".docx") or "word" in content_type or "openxmlformats" in content_type:
        return _extract_docx(content)
    elif fn.endswith(".xlsx") or "spreadsheet" in content_type or "excel" in content_type:
        return _extract_xlsx(content)
    elif fn.endswith(".csv") or "text/csv" in content_type or "csv" in content_type:
        return _extract_csv(content)
    elif fn.endswith((".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp", ".webp")) or "image/" in content_type:
        return _extract_image(content)
    elif fn.endswith(".txt") or "text/plain" in content_type:
        return content.decode("utf-8", errors="replace")
    else:
        # Best-effort: try UTF-8 text decode
        try:
            return content.decode("utf-8", errors="replace")
        except Exception:
            raise ValueError(f"Unsupported file type: {filename}")


def _ocr_image_bytes(png_bytes: bytes) -> str:
    """Run OCR on a single rasterized image. Returns '' (never raises) if the
    OCR stack (Pillow + pytesseract + the tesseract binary) is unavailable —
    callers must treat an empty result as 'could not read', not 'blank page'."""
    try:
        import pytesseract
        from PIL import Image

        img = Image.open(io.BytesIO(png_bytes))
        return pytesseract.image_to_string(img) or ""
    except ImportError:
        return ""
    except Exception:
        # pytesseract raises TesseractNotFoundError (subclass of EnvironmentError)
        # when the tesseract binary itself isn't installed on the host — treat
        # the same as "OCR unavailable" rather than failing the whole extraction.
        return ""


def _extract_pdf(content: bytes) -> str:
    """Extract text from PDF using PyMuPDF, falling back to OCR per-page for
    scanned/image-only pages (no native text layer)."""
    try:
        import fitz  # PyMuPDF

        parts = []
        ocr_pages = 0
        with fitz.open(stream=content, filetype="pdf") as doc:
            for i, page in enumerate(doc):
                text = page.get_text("text")
                if text.strip() and len(text.strip()) >= _OCR_FALLBACK_CHAR_THRESHOLD:
                    parts.append(f"[Page {i + 1}]\n{text.strip()}")
                    continue

                # No usable native text layer — rasterize the page and OCR it.
                # 2x zoom improves OCR accuracy on typical report-resolution scans.
                pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
                ocr_text = _ocr_image_bytes(pix.tobytes("png")).strip()
                if ocr_text:
                    ocr_pages += 1
                    parts.append(f"[Page {i + 1} (OCR)]\n{ocr_text}")
                elif text.strip():
                    # Sparse-but-nonzero native text and OCR unavailable/empty —
                    # keep the sparse native text rather than discarding it.
                    parts.append(f"[Page {i + 1}]\n{text.strip()}")

        full_text = "\n\n".join(parts)
        if ocr_pages:
            full_text = (
                f"[OCR NOTICE: {ocr_pages} page(s) in this document had no extractable "
                f"text layer and were machine-read via OCR — treat figures on those pages "
                f"with lower confidence than natively-extracted text.]\n\n{full_text}"
            )
        return _clean_text(full_text)
    except ImportError:
        raise RuntimeError("PyMuPDF not installed. Run: pip install pymupdf")


def _extract_image(content: bytes) -> str:
    """Extract text from a standalone image (photo of a certificate, scanned
    single-page document, etc.) via OCR. Raises if the OCR stack itself is
    unavailable — unlike the PDF path, there is no native-text fallback for a
    pure image, so a silent empty return would misrepresent 'not installed'
    as 'nothing found in this image'."""
    text = _ocr_image_bytes(content).strip()
    if not text:
        raise RuntimeError(
            "OCR unavailable or found no readable text in this image. "
            "Ensure pytesseract + the tesseract binary are installed on this host."
        )
    return _clean_text(f"[Image (OCR)]\n{text}")


def _extract_docx(content: bytes) -> str:
    """Extract text from DOCX using python-docx.

    Each paragraph/table-row is prefixed with a locator (¶N / TableT.RowR)
    so an AI review's citation can point at a real position in the source
    document instead of just "somewhere in this file" — needed for the
    Question Intelligence drawer's "View Source" feature.
    """
    try:
        from docx import Document

        doc = Document(io.BytesIO(content))
        lines = []

        para_idx = 0
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                para_idx += 1
                lines.append(f"[¶{para_idx}] {text}")

        for table_idx, table in enumerate(doc.tables, start=1):
            for row_idx, row in enumerate(table.rows, start=1):
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    lines.append(f"[Table {table_idx}.Row {row_idx}] " + " | ".join(cells))

        return _clean_text("\n\n".join(lines))
    except ImportError:
        raise RuntimeError("python-docx not installed. Run: pip install python-docx")


def _extract_xlsx(content: bytes) -> str:
    """Extract text from XLSX using openpyxl.

    Each row keeps its real 1-based spreadsheet row number as a locator
    (e.g. "[Sheet: Operating Assets, Row 12]") so citations can point at
    an actual row instead of just the filename.
    """
    try:
        import openpyxl

        wb = openpyxl.load_workbook(io.BytesIO(content), data_only=True)
        parts = []

        for sheet in wb.worksheets:
            parts.append(f"=== Sheet: {sheet.title} ===")
            for row_idx, row in enumerate(sheet.iter_rows(values_only=True), start=1):
                cells = [str(c) for c in row if c is not None and str(c).strip() not in ("", "None")]
                if cells:
                    parts.append(f"[Sheet: {sheet.title}, Row {row_idx}] " + "  |  ".join(cells))

        return _clean_text("\n".join(parts))
    except ImportError:
        raise RuntimeError("openpyxl not installed. Run: pip install openpyxl")


def _extract_csv(content: bytes) -> str:
    """Extract text from CSV — renders as labelled rows for the AI.

    Rows keep their real 1-based CSV row number (header = row 1) as a
    locator, matching how a human would refer to "row 12" when opening
    the same file in a spreadsheet app.
    """
    import csv

    try:
        text = content.decode("utf-8", errors="replace")
        reader = csv.reader(text.splitlines())
        rows   = list(reader)
        if not rows:
            return ""

        headers = rows[0] if rows else []
        parts   = []

        for row_idx, row in enumerate(rows[1:], start=2):
            if not any(cell.strip() for cell in row):
                continue
            # Pair each cell with its header for readable context
            pairs = []
            for i, cell in enumerate(row):
                if cell.strip():
                    label = headers[i].strip() if i < len(headers) else f"Column {i+1}"
                    pairs.append(f"{label}: {cell.strip()}")
            if pairs:
                parts.append(f"[Row {row_idx}] " + " | ".join(pairs))

        return _clean_text("\n".join(parts))
    except Exception as e:
        raise RuntimeError(f"CSV extraction failed: {e}")


def extract_from_url(url: str) -> dict:
    """
    Fetch a web page and extract its readable text as evidence (spec: website
    URLs, sustainability report landing pages, regulatory filing pages).

    Returns {"text": str, "accessed_at": iso8601 str, "final_url": str} on
    success. Raises RuntimeError on any fetch/parse failure — never returns
    an empty-but-successful result, so a caller can't mistake "the site
    blocked us" for "the page said nothing."

    Text is tagged "[Section N]" per top-level content block, matching the
    same locator-citation discipline as the file extractors above, so a
    citation is either a real section number or absent — never fabricated.
    """
    from datetime import datetime, timezone

    try:
        import httpx
    except ImportError:
        raise RuntimeError("httpx not installed. Run: pip install httpx")

    try:
        resp = httpx.get(
            url, timeout=15.0, follow_redirects=True,
            headers={"User-Agent": "Mozilla/5.0 (compatible; ClimactixEvidenceAgent/1.0)"},
        )
        resp.raise_for_status()
    except Exception as e:
        raise RuntimeError(f"Could not fetch URL: {e}")

    try:
        from bs4 import BeautifulSoup

        soup = BeautifulSoup(resp.text, "html.parser")
        for tag in soup(["script", "style", "noscript", "svg", "nav", "footer"]):
            tag.decompose()

        blocks = []
        for el in soup.find_all(["h1", "h2", "h3", "h4", "p", "li", "td", "th"]):
            block_text = el.get_text(" ", strip=True)
            if block_text and len(block_text) > 2:
                blocks.append(block_text)

        parts = [f"[Section {i + 1}] {b}" for i, b in enumerate(blocks)]
        text = _clean_text("\n\n".join(parts))
    except ImportError:
        raise RuntimeError("beautifulsoup4 not installed. Run: pip install beautifulsoup4")

    if not text.strip():
        raise RuntimeError("No readable text content found at this URL.")

    return {
        "text": text,
        "accessed_at": datetime.now(timezone.utc).isoformat(),
        "final_url": str(resp.url),
    }


def _clean_text(text: str) -> str:
    """Remove excessive whitespace and normalize text."""
    # Collapse 3+ blank lines to 2
    text = re.sub(r"\n{3,}", "\n\n", text)
    # Remove non-printable characters except newlines/tabs
    text = re.sub(r"[^\x09\x0A\x0D\x20-\x7E -￿]", " ", text)
    return text.strip()
