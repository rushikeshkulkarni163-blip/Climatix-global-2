"""
Climactix AI — FastAPI Backend
================================
Core endpoints:
  GET  /health                Health check
  POST /api/upload            Upload PDF / DOCX / XLSX → extracted text
  POST /api/generate          ESG text → full narratives + scores (9 modules)
  POST /api/export/pdf        Narratives → premium PDF report
  POST /api/export/docx       Narratives → DOCX report

REST aliases (spec-compliant):
  POST /upload-data           Alias for /api/upload
  POST /generate-insights     Alias for /api/generate
  GET  /investor-summary      Last cached investor output
  GET  /report                Last cached report metadata

Start:
  uvicorn main:app --reload --port 8000
"""

import io
import json
import os
from typing import Optional

from dotenv import load_dotenv
from fastapi import FastAPI, File, HTTPException, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import Response, JSONResponse, StreamingResponse
from pydantic import BaseModel

load_dotenv()

from services.extractor import extract_text
from services.ai_engine import generate_all_narratives
from services.scorer import calculate_scores
from services.assessment_report import stream_report
from services.simulation_brief import stream_brief
from services.greenwashing_scanner import scan_for_greenwashing, extract_data, validate_claims, extract_claims
from services.esg_framework_intelligence import run_intelligence_analysis, FRAMEWORK_REGISTRY, UNIFIED_ESG_MODEL, CROSS_FRAMEWORK_MAP
from routers.auth import router as auth_router
from routers.mfa import router as mfa_router
from routers.api_keys import router as api_keys_router
import database as db

# ── Template directory ─────────────────────────────────────────────────────────
_TEMPLATES_DIR = os.path.join(os.path.dirname(__file__), "templates")

# ── App init ──────────────────────────────────────────────────────────────────

app = FastAPI(
    title="Climactix AI API",
    description="ESG Intelligence Engine — powered by Claude · 9 AI modules",
    version="2.0.0",
)

# ── CORS ──────────────────────────────────────────────────────────────────────
# Credentials (cookies) require explicit origins — cannot use wildcard.
# Allow any localhost:PORT for dev; production origins set via CORS_ORIGINS env.
_env_origins = os.getenv("CORS_ORIGINS", "").strip()
_explicit_origins = [o.strip() for o in _env_origins.split(",") if o.strip()] if _env_origins else []

app.add_middleware(
    CORSMiddleware,
    allow_origins=_explicit_origins or ["http://localhost:3000", "http://localhost:8080", "http://localhost:5500", "http://127.0.0.1:5500"],
    allow_origin_regex=r"http://(localhost|127\.0\.0\.1)(:\d+)?",
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
    expose_headers=["Set-Cookie"],
)

# ── Auth + MFA + API key routers ──────────────────────────────────────────────
app.include_router(auth_router)
app.include_router(mfa_router)
app.include_router(api_keys_router)

# ── In-memory cache (last successful generation) ──────────────────────────────
_cache: dict = {}


# ── Request / Response models ─────────────────────────────────────────────────

class GenerateRequest(BaseModel):
    text: str
    company_name: Optional[str] = "The Company"
    report_year: Optional[str] = "2024"


class ExportRequest(BaseModel):
    narratives: dict
    scores: dict
    company_name: Optional[str] = "The Company"
    report_year: Optional[str] = "2024"


class AssessmentReportRequest(BaseModel):
    company_name: Optional[str] = "The Company"
    industry: Optional[str] = "Unknown Industry"
    report_year: Optional[str] = "FY 2025"
    eri_score: Optional[int] = 0
    tier_name: Optional[str] = "Unknown"
    tier_desc: Optional[str] = ""
    environmental_score: Optional[int] = 0
    social_score: Optional[int] = 0
    governance_score: Optional[int] = 0
    category_scores: Optional[list] = []
    framework_scores: Optional[list] = []
    key_disclosures: Optional[str] = ""


class SimulationBriefRequest(BaseModel):
    scenario_name: Optional[str] = "Delayed Transition"
    scenario_temp: Optional[str] = "2°C"
    scenario_desc: Optional[str] = ""
    year: Optional[int] = 2035
    portfolio_summary: Optional[dict] = {}
    assets: Optional[list] = []


# ── Health ─────────────────────────────────────────────────────────────────────

@app.get("/health")
def health():
    return {
        "status": "ok",
        "service": "Climactix AI API",
        "version": "2.0.0",
        "modules": 9,
    }


# ── Climate Data ───────────────────────────────────────────────────────────────

from services.climate_data import get_layer, LAYER_META, refresh_all as _refresh_climate

_VALID_LAYERS = list(LAYER_META.keys())


@app.on_event("startup")
async def _startup():
    """Initialize DB pool + auth schema, then warm the climate cache."""
    import asyncio, concurrent.futures
    # DB pool + schema migration (idempotent — safe to run on every startup)
    try:
        await db.init_pool()
    except Exception as exc:
        print(f"[WARN] DB init failed (auth features unavailable): {exc}")
    # Climate cache warm-up
    loop = asyncio.get_event_loop()
    with concurrent.futures.ThreadPoolExecutor() as pool:
        await loop.run_in_executor(pool, _refresh_climate)


@app.on_event("shutdown")
async def _shutdown():
    await db.close_pool()


@app.get("/api/climate-data")
def climate_data(layer: str = "temperature"):
    """
    Return a compact 72×36 normalised climate grid for the requested layer.

    Layers: temperature | co2 | sea_level | arctic_ice

    Response shape:
        {
          "layer": "temperature",
          "meta": { label, unit, source, description, legend, updated, data_source },
          "grid": { "rows": 36, "cols": 72, "values": [0.0–1.0 × 2592] }
        }
    """
    if layer not in _VALID_LAYERS:
        raise HTTPException(
            status_code=400,
            detail=f"Unknown layer '{layer}'. Valid values: {_VALID_LAYERS}",
        )
    try:
        data = get_layer(layer)
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Climate data error: {exc}")

    return JSONResponse(content=data)


@app.get("/api/climate-layers")
def climate_layers():
    """List all available climate layers with metadata."""
    return JSONResponse(content={
        "layers": [
            {
                "id":          k,
                "label":       v["label"],
                "unit":        v["unit"],
                "source":      v["source"],
                "description": v["description"],
                "color_scale": v["color_scale"],
                "legend":      v["legend"],
            }
            for k, v in LAYER_META.items()
        ]
    })


# ── Upload ─────────────────────────────────────────────────────────────────────

async def _handle_upload(file: UploadFile) -> dict:
    MAX_SIZE = 10 * 1024 * 1024  # 10 MB
    content = await file.read()
    if len(content) > MAX_SIZE:
        raise HTTPException(status_code=413, detail="File exceeds 10 MB limit.")
    try:
        text = extract_text(content, file.filename or "", file.content_type or "")
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except RuntimeError as e:
        raise HTTPException(status_code=500, detail=str(e))

    return {
        "success": True,
        "filename": file.filename,
        "word_count": len(text.split()),
        "char_count": len(text),
        "text": text,
    }


@app.post("/api/upload")
async def upload_file(file: UploadFile = File(...)):
    """Accept a PDF / DOCX / XLSX file and return extracted plain text."""
    return await _handle_upload(file)


@app.post("/upload-data")
async def upload_data(file: UploadFile = File(...)):
    """REST alias — same as /api/upload."""
    return await _handle_upload(file)


# ── Generate ───────────────────────────────────────────────────────────────────

def _handle_generate(req: GenerateRequest) -> dict:
    if not req.text.strip():
        raise HTTPException(status_code=400, detail="No ESG text provided.")
    try:
        narratives = generate_all_narratives(
            req.text, req.company_name or "The Company"
        )
    except RuntimeError as e:
        raise HTTPException(status_code=500, detail=str(e))
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"AI generation failed: {str(e)}")

    scores = calculate_scores(req.text, narratives)
    result = {
        "success": True,
        "company_name": req.company_name,
        "report_year": req.report_year,
        "narratives": narratives,
        "scores": scores,
    }
    # Cache for /investor-summary and /report
    _cache["last"] = result
    return result


@app.post("/api/generate")
def generate(req: GenerateRequest):
    """Generate all ESG narratives + scores from raw text (9 modules)."""
    return _handle_generate(req)


@app.post("/generate-insights")
def generate_insights(req: GenerateRequest):
    """REST alias — same as /api/generate."""
    return _handle_generate(req)


# ── Cached read endpoints ──────────────────────────────────────────────────────

@app.get("/investor-summary")
def investor_summary():
    """Return the investor_output section from the last successful generation."""
    if not _cache.get("last"):
        raise HTTPException(
            status_code=404,
            detail="No report generated yet. POST to /generate-insights first."
        )
    last = _cache["last"]
    return {
        "company_name": last.get("company_name"),
        "report_year": last.get("report_year"),
        "scores": last.get("scores"),
        "investor_output": last.get("narratives", {}).get("investor_output"),
        "investor_brief": last.get("narratives", {}).get("investor_brief"),
        "risk_flags": last.get("narratives", {}).get("risk_flags"),
    }


@app.get("/report")
def report_meta():
    """Return metadata for the last generated report (link to export endpoints)."""
    if not _cache.get("last"):
        raise HTTPException(
            status_code=404,
            detail="No report generated yet. POST to /generate-insights first."
        )
    last = _cache["last"]
    co   = (last.get("company_name") or "report").replace(" ", "-").lower()
    yr   = last.get("report_year") or "2024"
    return {
        "company_name": last.get("company_name"),
        "report_year": yr,
        "scores": last.get("scores"),
        "export_pdf":  f"/api/export/pdf",
        "export_docx": f"/api/export/docx",
        "filename": f"climactix-esg-{co}-{yr}",
    }


# ── HTML/WeasyPrint export ────────────────────────────────────────────────────

@app.post("/api/export/html-pdf")
def export_html_pdf(req: ExportRequest):
    """
    Premium PDF via Jinja2 HTML template + WeasyPrint.
    Falls back to a detailed error if WeasyPrint isn't installed.
    """
    try:
        from jinja2 import Environment, FileSystemLoader, select_autoescape
        from datetime import date

        env_j2 = Environment(
            loader=FileSystemLoader(_TEMPLATES_DIR),
            autoescape=select_autoescape(["html"]),
        )
        template = env_j2.get_template("esg_report.html")

        n  = req.narratives
        co = req.company_name or "The Company"
        yr = req.report_year  or "2024"

        report_title = (
            n.get("esg_summary", {}).get("report_title")
            or f"{co} ESG Intelligence Report {yr}"
        )

        html_str = template.render(
            company_name    = co,
            report_year     = yr,
            report_title    = report_title,
            narratives      = n,
            scores          = req.scores,
            investor_output = n.get("investor_output", {}),
            generated_date  = date.today().strftime("%-d %B %Y"),
        )
    except ImportError as e:
        raise HTTPException(
            status_code=500,
            detail=f"Missing dependency: {e}. Run: pip install jinja2 weasyprint"
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Template render failed: {e}")

    try:
        import weasyprint
        pdf_bytes = weasyprint.HTML(string=html_str).write_pdf()
    except ImportError:
        raise HTTPException(
            status_code=500,
            detail="WeasyPrint not installed. Run: pip install weasyprint"
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"WeasyPrint PDF failed: {e}")

    filename = f"climactix-esg-premium-{(req.company_name or 'report').replace(' ', '-').lower()}.pdf"
    return Response(
        content=pdf_bytes,
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# ── PDF export ─────────────────────────────────────────────────────────────────

@app.post("/api/export/pdf")
def export_pdf(req: ExportRequest):
    """Generate a premium styled PDF report from the narratives dict."""
    try:
        pdf_bytes = _build_pdf(req)
    except ImportError:
        raise HTTPException(
            status_code=500,
            detail="reportlab not installed. Run: pip install reportlab"
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"PDF generation failed: {str(e)}")

    filename = f"climactix-esg-report-{(req.company_name or 'report').replace(' ', '-').lower()}.pdf"
    return Response(
        content=pdf_bytes,
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# ── DOCX export ────────────────────────────────────────────────────────────────

@app.post("/api/export/docx")
def export_docx(req: ExportRequest):
    """Generate a DOCX report from the narratives dict."""
    try:
        docx_bytes = _build_docx(req)
    except ImportError:
        raise HTTPException(
            status_code=500,
            detail="python-docx not installed. Run: pip install python-docx"
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"DOCX generation failed: {str(e)}")

    filename = f"climactix-esg-report-{(req.company_name or 'report').replace(' ', '-').lower()}.docx"
    return Response(
        content=docx_bytes,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# ── ESG Risk Intelligence Scanner ────────────────────────────────────────────

class GreenwashingScanRequest(BaseModel):
    text: str
    company_name: Optional[str] = "The Company"


@app.post("/api/analyze-esg")
async def analyze_esg(file: UploadFile = File(...)):
    """
    Accept a PDF / DOCX / XLSX file, extract text, and run the full
    Greenwashing Risk Scanner pipeline. Returns a structured risk assessment.
    File size limit: 20 MB.
    """
    MAX_SIZE = 20 * 1024 * 1024  # 20 MB
    content = await file.read()

    if len(content) > MAX_SIZE:
        raise HTTPException(status_code=413, detail="File exceeds 20 MB limit.")

    try:
        text = extract_text(content, file.filename or "", file.content_type or "")
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except RuntimeError as e:
        raise HTTPException(status_code=500, detail=str(e))

    if not text.strip():
        raise HTTPException(status_code=400, detail="Could not extract readable text from the file.")

    company_name = (file.filename or "").rsplit(".", 1)[0].replace("-", " ").replace("_", " ").title() or "The Company"

    try:
        result = scan_for_greenwashing(text, company_name)
        # Enrich with intelligence layer (jurisdiction, multi-framework coverage, integrity score)
        extracted = result.get("data_extracted", {})
        flags     = result.get("risk_flags", [])
        intel     = run_intelligence_analysis(text, extracted, flags, company_name)
        result["intelligence"] = intel
    except RuntimeError as e:
        raise HTTPException(status_code=500, detail=str(e))
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Scanner failed: {str(e)}")

    return JSONResponse(content={"success": True, **result})


@app.post("/api/analyze-esg/export-pdf")
def export_greenwashing_pdf(result: dict):
    """Generate a downloadable investor-grade PDF from a greenwashing scan result."""
    try:
        pdf_bytes = _build_greenwashing_pdf(result)
    except ImportError:
        raise HTTPException(status_code=500, detail="reportlab not installed.")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"PDF generation failed: {str(e)}")

    company = (result.get("company_name") or "report").replace(" ", "-").lower()
    filename = f"climactix-greenwashing-scan-{company}.pdf"
    return Response(
        content=pdf_bytes,
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@app.get("/api/esg-intelligence/frameworks")
def list_frameworks():
    """
    Return the full ESG Framework Registry — metadata for all 16 supported
    standards (GRI, TCFD, ISSB S1/S2, CSRD, ESRS, BRSR, SFDR, SEC,
    TNFD, GRESB, ISO 14001, PCAF, GHG Protocol, SBTi).
    """
    return JSONResponse(content={
        "count": len(FRAMEWORK_REGISTRY),
        "frameworks": {
            fid: {
                "full_name": meta["full_name"],
                "type": meta["type"],
                "focus": meta["focus"],
                "jurisdiction": meta["jurisdiction"],
                "target_users": meta["target_users"],
                "disclosure_areas": meta["disclosure_areas"],
                "mandatory_in": meta["mandatory_in"],
            }
            for fid, meta in FRAMEWORK_REGISTRY.items()
        },
    })


@app.get("/api/esg-intelligence/model")
def unified_model():
    """
    Return the Unified ESG Data Model — canonical concepts with cross-framework
    mappings (which frameworks require each disclosure point).
    """
    return JSONResponse(content={
        "concept_count": len(UNIFIED_ESG_MODEL),
        "categories": sorted({c["category"] for c in UNIFIED_ESG_MODEL.values()}),
        "concepts": {
            cid: {
                "category": c["category"],
                "sub_category": c["sub_category"],
                "concept": c["concept"],
                "metric_type": c["metric_type"],
                "framework_mappings": {
                    fw: {"ref": info["ref"], "weight": info["weight"]}
                    for fw, info in c["framework_mappings"].items()
                },
                "scoring_weight": c["scoring_weight"],
            }
            for cid, c in UNIFIED_ESG_MODEL.items()
        },
    })


@app.get("/api/esg-intelligence/cross-framework-map")
def cross_framework_map():
    """
    Return the cross-framework equivalence table — one disclosure satisfies
    multiple frameworks.  Shows interoperability across GRI, TCFD, ISSB, CSRD,
    BRSR, SEC, CDP, SBTi, GHG Protocol.
    """
    return JSONResponse(content={
        "description": (
            "Maps each ESG concept to the framework references it satisfies. "
            "A single disclosure on Scope 1 emissions, for example, satisfies "
            "GRI 305-1, TCFD Metrics, ISSB S2 para 29, CSRD ESRS E1-6, BRSR P9, "
            "SEC Climate Rule, CDP C6.1, and GHG Protocol simultaneously."
        ),
        "cross_framework_map": CROSS_FRAMEWORK_MAP,
    })


@app.post("/api/esg-intelligence/analyze")
async def esg_intelligence_analyze(file: UploadFile = File(...)):
    """
    Upload an ESG report → run full Framework Intelligence pipeline.
    Returns:
      - jurisdiction detection (India → BRSR, EU → CSRD, etc.)
      - coverage matrix against unified ESG model
      - per-framework compliance status (Aligned / Partial / Missing)
      - ESG Integrity Score (0–100, higher = better)
      - Greenwashing Risk Score (0–100, inverse)
      - Critical gaps with cross-framework refs
    """
    MAX_SIZE = 20 * 1024 * 1024
    content = await file.read()
    if len(content) > MAX_SIZE:
        raise HTTPException(status_code=413, detail="File exceeds 20 MB limit.")

    try:
        text = extract_text(content, file.filename or "", file.content_type or "")
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except RuntimeError as e:
        raise HTTPException(status_code=500, detail=str(e))

    if not text.strip():
        raise HTTPException(status_code=400, detail="No readable text extracted.")

    company_name = (file.filename or "").rsplit(".", 1)[0].replace("-", " ").replace("_", " ").title() or "The Company"

    try:
        extracted = extract_data(text)
        claims    = extract_claims(text)
        flags     = validate_claims(claims, extracted, text)
        intel     = run_intelligence_analysis(text, extracted, flags, company_name)
    except RuntimeError as e:
        raise HTTPException(status_code=500, detail=str(e))
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Intelligence analysis failed: {str(e)}")

    return JSONResponse(content={"success": True, **intel})


def _build_greenwashing_pdf(result: dict) -> bytes:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, HRFlowable, Table, TableStyle,
    )
    from datetime import date

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        topMargin=20 * mm, bottomMargin=20 * mm,
        leftMargin=22 * mm, rightMargin=22 * mm,
    )

    # Palette
    dark     = colors.HexColor("#050A14")
    forest   = colors.HexColor("#2d5a3d")
    emerald  = colors.HexColor("#10B981")
    amber    = colors.HexColor("#F59E0B")
    danger   = colors.HexColor("#EF4444")
    ink      = colors.HexColor("#1a1a2e")
    muted    = colors.HexColor("#6B7280")
    white    = colors.white

    risk_score = result.get("risk_score", 0)
    risk_level = result.get("risk_level", "Unknown")
    score_color = (
        emerald if risk_level == "Low"
        else amber if risk_level == "Medium"
        else danger
    )

    h1  = ParagraphStyle("H1",  fontName="Helvetica-Bold",   fontSize=20, textColor=dark,   spaceAfter=4)
    h2  = ParagraphStyle("H2",  fontName="Helvetica-Bold",   fontSize=13, textColor=forest, spaceBefore=12, spaceAfter=4)
    h3  = ParagraphStyle("H3",  fontName="Helvetica-Bold",   fontSize=10, textColor=ink,    spaceBefore=8,  spaceAfter=3)
    bod = ParagraphStyle("Bod", fontName="Helvetica",        fontSize=9,  textColor=ink,    leading=14,     spaceAfter=6)
    met = ParagraphStyle("Met", fontName="Helvetica",        fontSize=8,  textColor=muted,  spaceAfter=3)
    sev_high = ParagraphStyle("SH", fontName="Helvetica-Bold", fontSize=9, textColor=danger)
    sev_med  = ParagraphStyle("SM", fontName="Helvetica-Bold", fontSize=9, textColor=amber)
    sev_low  = ParagraphStyle("SL", fontName="Helvetica-Bold", fontSize=9, textColor=emerald)

    story = []
    co = result.get("company_name", "The Company")

    # Header
    story.append(Paragraph("CLIMACTIX AI", ParagraphStyle(
        "Brand", fontName="Helvetica-Bold", fontSize=10, textColor=forest, spaceAfter=2,
    )))
    story.append(Paragraph("ESG Risk Intelligence · Investor-Grade ESG Disclosure Audit", met))
    story.append(Spacer(1, 3 * mm))
    story.append(HRFlowable(width="100%", thickness=2, color=forest))
    story.append(Spacer(1, 4 * mm))
    story.append(Paragraph(f"Greenwashing Risk Report — {co}", h1))
    story.append(Paragraph(f"Generated {date.today().strftime('%-d %B %Y')} · Climactix AI", met))
    story.append(Spacer(1, 5 * mm))

    # Score summary table
    bd = result.get("score_breakdown", {})
    score_data = [
        ["Risk Score", "Risk Level", "Mismatch Score", "Missing Disclosures", "Framework Gaps"],
        [
            f"{risk_score}/100",
            risk_level,
            f"{bd.get('narrative_data_mismatch',{}).get('score',0)}/100",
            f"{bd.get('missing_disclosures',{}).get('items_missing',0)} items",
            f"{100 - result.get('framework_coverage_pct', 0)}% unmet",
        ],
    ]
    t = Table(score_data, colWidths=[30*mm, 30*mm, 38*mm, 38*mm, 30*mm])
    t.setStyle(TableStyle([
        ("BACKGROUND",    (0, 0), (-1, 0), forest),
        ("TEXTCOLOR",     (0, 0), (-1, 0), white),
        ("FONTNAME",      (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE",      (0, 0), (-1, 0), 8),
        ("FONTNAME",      (0, 1), (-1, 1), "Helvetica-Bold"),
        ("FONTSIZE",      (0, 1), (0, 1),  18),
        ("TEXTCOLOR",     (0, 1), (0, 1),  score_color),
        ("TEXTCOLOR",     (1, 1), (1, 1),  score_color),
        ("FONTSIZE",      (1, 1), (-1, 1), 11),
        ("ALIGN",         (0, 0), (-1, -1), "CENTER"),
        ("VALIGN",        (0, 0), (-1, -1), "MIDDLE"),
        ("ROWBACKGROUNDS",(0, 1), (-1, -1), [colors.HexColor("#F4F9F8")]),
        ("GRID",          (0, 0), (-1, -1), 0.5, white),
        ("TOPPADDING",    (0, 0), (-1, -1), 7),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
    ]))
    story.append(t)
    story.append(Spacer(1, 7 * mm))

    def _hr():
        story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#d0d0c0")))
        story.append(Spacer(1, 2 * mm))

    # Risk Flags
    flags = result.get("risk_flags", [])
    if flags:
        _hr()
        story.append(Paragraph("Risk Flags Identified", h2))
        sev_map = {"High": sev_high, "Medium": sev_med, "Low": sev_low}
        for f in flags:
            sty = sev_map.get(f.get("severity", "Medium"), sev_med)
            story.append(Paragraph(f"[{f.get('severity','').upper()}] {f.get('title','')}", sty))
            story.append(Paragraph(f["description"], bod))
            story.append(Paragraph(f"Framework: {f.get('framework_ref','')}", met))
            if f.get("claim_excerpt"):
                story.append(Paragraph(f"Excerpt: \"{f['claim_excerpt']}\"", met))
            story.append(Spacer(1, 3 * mm))

    # Missing Disclosures
    missing = result.get("missing_disclosures", [])
    if missing:
        _hr()
        story.append(Paragraph("Missing Framework Disclosures", h2))
        for item in missing:
            story.append(Paragraph(f"✗  {item.get('requirement','')} [{item.get('framework','')}]", h3))
            story.append(Paragraph(item.get("description", ""), bod))

    # Recommendations
    recs = result.get("recommendations", [])
    if recs:
        _hr()
        story.append(Paragraph("Remediation Recommendations", h2))
        for r in recs:
            pri = r.get("priority", "Medium")
            sty = sev_map.get(pri, sev_med)
            story.append(Paragraph(f"[{pri.upper()}] {r.get('title','')}", sty))
            story.append(Paragraph(r.get("action", ""), bod))
            story.append(Paragraph(f"Framework: {r.get('framework','')}  ·  Impact: {r.get('impact','')}", met))
            story.append(Spacer(1, 3 * mm))

    # Footer
    story.append(Spacer(1, 8 * mm))
    story.append(HRFlowable(width="100%", thickness=1, color=forest))
    story.append(Spacer(1, 2 * mm))
    story.append(Paragraph(
        "Generated by Climactix AI ESG Risk Intelligence · climactixglobal.com · Powered by Claude",
        ParagraphStyle("Footer", fontName="Helvetica", fontSize=7, textColor=muted, alignment=1),
    ))

    doc.build(story)
    return buf.getvalue()


# ── PDF builder ────────────────────────────────────────────────────────────────

def _build_pdf(req: ExportRequest) -> bytes:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, HRFlowable, Table, TableStyle, KeepTogether
    )

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        topMargin=20 * mm, bottomMargin=20 * mm,
        leftMargin=22 * mm, rightMargin=22 * mm,
    )

    # Colour palette
    forest  = colors.HexColor("#2d5a3d")
    env_col = colors.HexColor("#2d8a5e")
    soc_col = colors.HexColor("#1a6b8a")
    gov_col = colors.HexColor("#b8763a")
    ink     = colors.HexColor("#1a1a14")
    muted   = colors.HexColor("#7a7a68")
    cream   = colors.HexColor("#f5f2eb")
    danger  = colors.HexColor("#cc3311")
    warn    = colors.HexColor("#dd8800")

    styles = getSampleStyleSheet()
    h1   = ParagraphStyle("H1",   fontName="Helvetica-Bold",   fontSize=22, textColor=forest, spaceAfter=6)
    h2   = ParagraphStyle("H2",   fontName="Helvetica-Bold",   fontSize=14, textColor=forest, spaceBefore=14, spaceAfter=4)
    h3   = ParagraphStyle("H3",   fontName="Helvetica-Bold",   fontSize=11, textColor=ink,    spaceBefore=10, spaceAfter=3)
    body = ParagraphStyle("Body", fontName="Helvetica",        fontSize=10, textColor=ink,    leading=15, spaceAfter=8)
    meta = ParagraphStyle("Meta", fontName="Helvetica",        fontSize=9,  textColor=muted,  spaceAfter=4)
    quot = ParagraphStyle("Quot", fontName="Helvetica-Oblique",fontSize=12, textColor=forest, spaceAfter=8, leftIndent=10)
    rec  = ParagraphStyle("Rec",  fontName="Helvetica-Bold",   fontSize=10, textColor=forest, spaceAfter=4)

    story = []
    n  = req.narratives
    co = req.company_name
    yr = req.report_year
    sc = req.scores

    # ── Cover header ──
    story.append(Paragraph("CLIMACTIX AI", ParagraphStyle(
        "Brand", fontName="Helvetica-Bold", fontSize=11, textColor=forest, spaceAfter=2,
    )))
    story.append(Paragraph("ESG Intelligence Report · Full Analysis", meta))
    story.append(Spacer(1, 4 * mm))
    story.append(HRFlowable(width="100%", thickness=2, color=forest))
    story.append(Spacer(1, 4 * mm))

    title = (n.get("esg_summary", {}).get("report_title") or f"{co} ESG Report {yr}")
    story.append(Paragraph(title, h1))
    story.append(Paragraph(f"Prepared by Climactix AI · {co} · {yr}", meta))
    story.append(Spacer(1, 6 * mm))

    # ── ESG Grade + Main Scores ──
    inv_out = n.get("investor_output", {})
    grade   = inv_out.get("esg_grade", "B")
    grade_data = [
        ["ESG Grade", "Overall Score", "Environmental", "Social", "Governance"],
        [
            grade,
            f"{sc.get('esg_score', '--')}/100",
            f"{sc.get('environmental_score', '--')}/100",
            f"{sc.get('social_score', '--')}/100",
            f"{sc.get('governance_score', '--')}/100",
        ],
    ]
    t = Table(grade_data, colWidths=[32*mm, 34*mm, 34*mm, 32*mm, 34*mm])
    t.setStyle(TableStyle([
        ("BACKGROUND",    (0, 0), (-1, 0), forest),
        ("TEXTCOLOR",     (0, 0), (-1, 0), colors.white),
        ("FONTNAME",      (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE",      (0, 0), (-1, 0), 8),
        ("FONTNAME",      (0, 1), (-1, 1), "Helvetica-Bold"),
        ("FONTSIZE",      (0, 1), (0, 1),  20),
        ("FONTSIZE",      (1, 1), (-1, 1), 15),
        ("TEXTCOLOR",     (0, 1), (0, 1),  forest),
        ("TEXTCOLOR",     (1, 1), (1, 1),  forest),
        ("TEXTCOLOR",     (2, 1), (2, 1),  env_col),
        ("TEXTCOLOR",     (3, 1), (3, 1),  soc_col),
        ("TEXTCOLOR",     (4, 1), (4, 1),  gov_col),
        ("ALIGN",         (0, 0), (-1, -1), "CENTER"),
        ("VALIGN",        (0, 0), (-1, -1), "MIDDLE"),
        ("ROWBACKGROUNDS",(0, 1), (-1, -1), [cream]),
        ("GRID",          (0, 0), (-1, -1), 0.5, colors.white),
        ("TOPPADDING",    (0, 0), (-1, -1), 8),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
    ]))
    story.append(t)
    story.append(Spacer(1, 8 * mm))

    def _section(title_text: str, content: str, col=None):
        story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#d0d0c0")))
        story.append(Spacer(1, 2 * mm))
        h = ParagraphStyle("SH", fontName="Helvetica-Bold", fontSize=14,
                           textColor=col or forest, spaceBefore=14, spaceAfter=4)
        story.append(Paragraph(title_text, h))
        for para in content.split("\n\n"):
            para = para.strip()
            if para:
                story.append(Paragraph(para, body))
        story.append(Spacer(1, 4 * mm))

    # ── 1. Investor Brief ──
    ib = n.get("investor_brief", {})
    story.append(Paragraph("1. Investor Brief", h2))
    if ib.get("headline"):
        story.append(Paragraph(f'"{ib["headline"]}"', quot))
    if ib.get("executive_summary"):
        story.append(Paragraph(ib["executive_summary"], body))
    if ib.get("key_metrics"):
        story.append(Paragraph("Key Metrics", h3))
        for m in ib["key_metrics"]:
            story.append(Paragraph(f"• {m}", body))
    if ib.get("risks"):
        story.append(Paragraph("Risks", h3))
        for r in ib["risks"]:
            story.append(Paragraph(f"→ {r}", body))
    if ib.get("opportunities"):
        story.append(Paragraph("Opportunities", h3))
        for o in ib["opportunities"]:
            story.append(Paragraph(f"✓ {o}", body))
    if ib.get("recommendation"):
        story.append(Paragraph(f"Recommendation: {ib['recommendation']}", rec))
    story.append(Spacer(1, 6 * mm))

    # ── 2. Investor Output ──
    if inv_out:
        story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#d0d0c0")))
        story.append(Spacer(1, 2 * mm))
        story.append(Paragraph("2. Investor Output — Sustainability Positioning", h2))
        if inv_out.get("sustainability_positioning"):
            story.append(Paragraph(inv_out["sustainability_positioning"], body))
        if inv_out.get("investment_thesis"):
            story.append(Paragraph("Investment Thesis", h3))
            story.append(Paragraph(inv_out["investment_thesis"], body))
        for pillar, label, col in [
            ("e_score_rationale", "Environmental Pillar", env_col),
            ("s_score_rationale", "Social Pillar",        soc_col),
            ("g_score_rationale", "Governance Pillar",    gov_col),
        ]:
            if inv_out.get(pillar):
                p = ParagraphStyle(f"P{pillar}", fontName="Helvetica-Bold", fontSize=10,
                                   textColor=col, spaceAfter=4)
                story.append(Paragraph(label, p))
                story.append(Paragraph(inv_out[pillar], body))
        if inv_out.get("growth_opportunities"):
            story.append(Paragraph("Growth Opportunities", h3))
            for o in inv_out["growth_opportunities"]:
                story.append(Paragraph(f"✓ {o}", body))
        story.append(Spacer(1, 6 * mm))

    # ── 3. Risk Flags ──
    risk_flags = n.get("risk_flags", [])
    if risk_flags:
        story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#d0d0c0")))
        story.append(Spacer(1, 2 * mm))
        story.append(Paragraph("3. ESG Risk Flags", h2))
        severity_color = {"High": danger, "Medium": warn, "Low": muted}
        for flag in risk_flags:
            sev  = flag.get("severity", "Medium")
            col  = severity_color.get(sev, muted)
            cat  = flag.get("category", "")
            pill = ParagraphStyle(f"Pill{sev}", fontName="Helvetica-Bold", fontSize=9,
                                  textColor=col, spaceAfter=2)
            story.append(Paragraph(f"[{sev.upper()}] {flag.get('title','')} · {cat}", pill))
            story.append(Paragraph(flag.get("description", ""), body))
            if flag.get("mitigation"):
                story.append(Paragraph(f"Mitigation: {flag['mitigation']}", meta))
            story.append(Spacer(1, 3 * mm))

    # ── 4. ESG Summary ──
    es = n.get("esg_summary", {})
    if es.get("executive_summary"):
        _section("4. ESG Executive Summary", es["executive_summary"])

    # ── 5. E/S/G Pillar Summaries ──
    for key, label, col in [
        ("environmental", "5a. Environmental Performance", env_col),
        ("social",        "5b. Social Performance",        soc_col),
        ("governance",    "5c. Governance & Accountability", gov_col),
    ]:
        if es.get(key):
            _section(label, es[key], col)

    if es.get("forward_guidance"):
        _section("5d. Forward Guidance & Targets", es["forward_guidance"])

    # ── 6. Insights ──
    insights = n.get("insights", {})
    if insights:
        story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#d0d0c0")))
        story.append(Spacer(1, 2 * mm))
        story.append(Paragraph("6. ESG Insights", h2))
        improving = insights.get("improving", [])
        if improving:
            story.append(Paragraph("Improving", h3))
            for item in improving:
                story.append(Paragraph(f"↑ {item.get('title','')} [{item.get('category','')}]", rec))
                story.append(Paragraph(item.get("detail", ""), body))
        attention = insights.get("needs_attention", [])
        if attention:
            story.append(Paragraph("Needs Attention", h3))
            for item in attention:
                story.append(Paragraph(f"! {item.get('title','')} [{item.get('urgency','')}]", rec))
                story.append(Paragraph(item.get("detail", ""), body))
        if insights.get("strategic_narrative"):
            story.append(Spacer(1, 3 * mm))
            story.append(Paragraph(insights["strategic_narrative"], body))
        story.append(Spacer(1, 6 * mm))

    # ── 7. Marketing Narrative ──
    mn = n.get("marketing_narrative", {})
    if mn.get("story"):
        _section("7. Marketing Narrative", mn["story"])
        if mn.get("tagline"):
            story.append(Paragraph(f'Tagline: "{mn["tagline"]}"', rec))
        story.append(Spacer(1, 4 * mm))

    # ── 8. Social Media ──
    sm = n.get("social_media", {})
    story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#d0d0c0")))
    story.append(Spacer(1, 2 * mm))
    story.append(Paragraph("8. Social Media Content", h2))
    if sm.get("linkedin"):
        story.append(Paragraph("LinkedIn Post", h3))
        story.append(Paragraph(sm["linkedin"], body))
    if sm.get("instagram"):
        story.append(Paragraph("Instagram Caption", h3))
        story.append(Paragraph(sm["instagram"], body))
    story.append(Spacer(1, 6 * mm))

    # ── 9. SDG Mapping ──
    sdg = n.get("sdg_mapping", {})
    story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#d0d0c0")))
    story.append(Spacer(1, 2 * mm))
    story.append(Paragraph("9. UN SDG Alignment", h2))
    if sdg.get("alignment_summary"):
        story.append(Paragraph(sdg["alignment_summary"], body))
    for sdg_item in sdg.get("top_sdgs", []):
        story.append(Paragraph(
            f'{sdg_item.get("icon", "")} SDG {sdg_item.get("number")} — {sdg_item.get("name")} '
            f'({sdg_item.get("relevance_label", "")} / {sdg_item.get("relevance_score", "")})',
            h3,
        ))
        story.append(Paragraph(sdg_item.get("explanation", ""), body))
        if sdg_item.get("evidence"):
            story.append(Paragraph(f"Evidence: {sdg_item['evidence']}", meta))

    # ── Footer ──
    story.append(Spacer(1, 10 * mm))
    story.append(HRFlowable(width="100%", thickness=1, color=forest))
    story.append(Spacer(1, 2 * mm))
    story.append(Paragraph(
        "Generated by Climactix AI · climactixglobal.com · Powered by Claude",
        ParagraphStyle("Footer", fontName="Helvetica", fontSize=8,
                       textColor=muted, alignment=1),
    ))

    doc.build(story)
    return buf.getvalue()


# ── DOCX builder ───────────────────────────────────────────────────────────────

def _build_docx(req: ExportRequest) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    forest_rgb  = RGBColor(0x2D, 0x5A, 0x3D)
    env_rgb     = RGBColor(0x2D, 0x8A, 0x5E)
    soc_rgb     = RGBColor(0x1A, 0x6B, 0x8A)
    gov_rgb     = RGBColor(0xB8, 0x76, 0x3A)
    ink_rgb     = RGBColor(0x1A, 0x1A, 0x14)
    danger_rgb  = RGBColor(0xCC, 0x33, 0x11)

    def heading(text, level=1, color=None):
        p = doc.add_heading(text, level=level)
        if color:
            for run in p.runs:
                run.font.color.rgb = color
        return p

    def body_para(text: str, bold=False, color=None):
        if not text:
            return
        for block in text.split("\n\n"):
            block = block.strip()
            if block:
                p = doc.add_paragraph(block)
                p.paragraph_format.space_after = Pt(8)
                if bold or color:
                    for run in p.runs:
                        if bold: run.bold = True
                        if color: run.font.color.rgb = color

    n  = req.narratives
    co = req.company_name
    yr = req.report_year
    sc = req.scores

    # Title
    title = (n.get("esg_summary", {}).get("report_title") or f"{co} ESG Report {yr}")
    heading(title, level=1, color=forest_rgb)
    doc.add_paragraph(f"Climactix AI · {co} · {yr}")
    doc.add_paragraph()

    # Scores table
    heading("ESG Scores", level=2, color=forest_rgb)
    tbl = doc.add_table(rows=2, cols=5)
    tbl.style = "Table Grid"
    for i, h in enumerate(["ESG Grade", "Overall Score", "Environmental", "Social", "Governance"]):
        tbl.rows[0].cells[i].text = h
    inv_out = n.get("investor_output", {})
    for i, v in enumerate([
        inv_out.get("esg_grade", "--"),
        f"{sc.get('esg_score', '--')}/100",
        f"{sc.get('environmental_score', '--')}/100",
        f"{sc.get('social_score', '--')}/100",
        f"{sc.get('governance_score', '--')}/100",
    ]):
        tbl.rows[1].cells[i].text = v
    doc.add_paragraph()

    # 1. Investor Brief
    ib = n.get("investor_brief", {})
    heading("1. Investor Brief", level=2, color=forest_rgb)
    if ib.get("headline"):
        p = doc.add_paragraph(f'"{ib["headline"]}"')
        p.runs[0].bold = True
    body_para(ib.get("executive_summary", ""))
    if ib.get("key_metrics"):
        heading("Key Metrics", level=3, color=forest_rgb)
        for m in ib["key_metrics"]:
            doc.add_paragraph(m, style="List Bullet")
    if ib.get("risks"):
        heading("Risks", level=3, color=danger_rgb)
        for r in ib["risks"]:
            doc.add_paragraph(r, style="List Bullet")
    if ib.get("opportunities"):
        heading("Opportunities", level=3, color=forest_rgb)
        for o in ib["opportunities"]:
            doc.add_paragraph(o, style="List Bullet")
    if ib.get("recommendation"):
        p = doc.add_paragraph(f"Recommendation: {ib['recommendation']}")
        p.runs[0].bold = True
    doc.add_paragraph()

    # 2. Investor Output
    if inv_out:
        heading("2. Investor Output", level=2, color=forest_rgb)
        body_para(inv_out.get("sustainability_positioning", ""))
        if inv_out.get("investment_thesis"):
            heading("Investment Thesis", level=3, color=forest_rgb)
            body_para(inv_out.get("investment_thesis", ""))
        if inv_out.get("growth_opportunities"):
            heading("Growth Opportunities", level=3, color=forest_rgb)
            for o in inv_out["growth_opportunities"]:
                doc.add_paragraph(o, style="List Bullet")
        doc.add_paragraph()

    # 3. Risk Flags
    risk_flags = n.get("risk_flags", [])
    if risk_flags:
        heading("3. ESG Risk Flags", level=2, color=danger_rgb)
        for flag in risk_flags:
            p = doc.add_paragraph(f"[{flag.get('severity','')}] {flag.get('title','')} · {flag.get('category','')}")
            p.runs[0].bold = True
            if flag.get("severity") == "High":
                p.runs[0].font.color.rgb = danger_rgb
            body_para(flag.get("description", ""))
        doc.add_paragraph()

    # 4. ESG Summary
    es = n.get("esg_summary", {})
    heading("4. ESG Executive Summary", level=2, color=forest_rgb)
    body_para(es.get("executive_summary", ""))
    for key, label, col in [
        ("environmental", "Environmental", env_rgb),
        ("social",        "Social",        soc_rgb),
        ("governance",    "Governance",    gov_rgb),
    ]:
        if es.get(key):
            heading(label, level=3, color=col)
            body_para(es[key])
    if es.get("forward_guidance"):
        heading("Forward Guidance", level=3, color=forest_rgb)
        body_para(es["forward_guidance"])
    doc.add_paragraph()

    # 5. Insights
    insights = n.get("insights", {})
    if insights:
        heading("5. ESG Insights", level=2, color=forest_rgb)
        for item in insights.get("improving", []):
            p = doc.add_paragraph(f"↑ {item.get('title','')} [{item.get('category','')}]")
            p.runs[0].bold = True
            body_para(item.get("detail", ""))
        for item in insights.get("needs_attention", []):
            p = doc.add_paragraph(f"! {item.get('title','')} [{item.get('urgency','')}]")
            p.runs[0].bold = True
            body_para(item.get("detail", ""))
        body_para(insights.get("strategic_narrative", ""))
        doc.add_paragraph()

    # 6. Marketing Narrative
    mn = n.get("marketing_narrative", {})
    heading("6. Marketing Narrative", level=2, color=forest_rgb)
    if mn.get("headline"):
        p = doc.add_paragraph(mn["headline"])
        p.runs[0].bold = True
    body_para(mn.get("story", ""))
    if mn.get("tagline"):
        p = doc.add_paragraph(f'Tagline: "{mn["tagline"]}"')
        p.runs[0].italic = True
    doc.add_paragraph()

    # 7. Social Media
    sm = n.get("social_media", {})
    heading("7. Social Media Content", level=2, color=forest_rgb)
    if sm.get("linkedin"):
        heading("LinkedIn Post", level=3, color=forest_rgb)
        body_para(sm["linkedin"])
    if sm.get("twitter_thread"):
        heading("X / Twitter Thread", level=3, color=forest_rgb)
        for tweet in sm["twitter_thread"]:
            doc.add_paragraph(tweet, style="List Bullet")
    if sm.get("instagram"):
        heading("Instagram Caption", level=3, color=forest_rgb)
        body_para(sm["instagram"])
    doc.add_paragraph()

    # 8. SDG Mapping
    sdg = n.get("sdg_mapping", {})
    heading("8. UN SDG Alignment", level=2, color=forest_rgb)
    body_para(sdg.get("alignment_summary", ""))
    for sdg_item in sdg.get("top_sdgs", []):
        heading(
            f'{sdg_item.get("icon", "")} SDG {sdg_item.get("number")} — {sdg_item.get("name")}',
            level=3, color=forest_rgb,
        )
        body_para(sdg_item.get("explanation", ""))
        if sdg_item.get("evidence"):
            doc.add_paragraph(f"Evidence: {sdg_item['evidence']}", style="List Bullet")
    doc.add_paragraph()

    # Footer
    doc.add_paragraph("─" * 60)
    footer_p = doc.add_paragraph("Generated by Climactix AI · climactixglobal.com · Powered by Claude")
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if footer_p.runs:
        footer_p.runs[0].font.size = Pt(8)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── OpenAI Assessment Report (streaming) ──────────────────────────────────────

@app.post("/api/assessment/report")
def assessment_report(req: AssessmentReportRequest):
    """
    Stream a full institutional ESG report for a completed assessment.
    Uses OpenAI GPT-4o with Server-Sent Events so the frontend
    can render the report in real-time as it is generated.
    """
    if not os.getenv("OPENAI_API_KEY"):
        raise HTTPException(
            status_code=503,
            detail="OpenAI API key not configured. Set OPENAI_API_KEY in backend/.env"
        )

    payload = req.model_dump()

    return StreamingResponse(
        stream_report(payload),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "X-Accel-Buffering": "no",
        },
    )


# ── Climate Risk Simulation Intelligence Brief (streaming) ────────────────────

@app.post("/api/simulation/risk-brief")
def simulation_risk_brief(req: SimulationBriefRequest):
    """
    Stream a GPT-4o climate risk intelligence brief for the current
    simulation portfolio snapshot.  Uses Server-Sent Events so the frontend
    renders the brief in real-time.
    """
    if not os.getenv("OPENAI_API_KEY"):
        raise HTTPException(
            status_code=503,
            detail="OpenAI API key not configured. Set OPENAI_API_KEY in backend/.env"
        )

    payload = req.model_dump()

    return StreamingResponse(
        stream_brief(payload),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "X-Accel-Buffering": "no",
        },
    )
